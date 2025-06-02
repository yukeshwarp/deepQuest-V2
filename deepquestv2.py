import streamlit as st
from dotenv import load_dotenv
from writer import report_writer
from planner import plan_research, replanner
from stepexecutor import execute_step
from io import BytesIO
from docx import Document
from bs4 import BeautifulSoup
import markdown as md
import logging

load_dotenv()

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s"
)

st.title("deepQuest v2")

# Set up the sidebar
st.sidebar.title("Research Steps")
@st.cache_data

def agentic_brain(query, max_replan_rounds=3, max_total_steps=20):
    try:
        steps = plan_research(query)
        sidebar_steps = st.sidebar.empty()
        sidebar_steps.markdown(
            "\n".join([f"{idx+1}. {step}" for idx, step in enumerate(steps)])
        )

        context = ""
        completed_steps = []
        i = 0
        replan_rounds = 0
        replan_limit_reached = False
        max_steps_warning_shown = False
        while i < len(steps):
            if len(steps) > max_total_steps and not max_steps_warning_shown:
                st.warning(
                    "Maximum total steps reached. No further replanning will be done, but all planned steps will be executed."
                )
                max_steps_warning_shown = True
                replan_limit_reached = True  # Prevent further replanning

            step = steps[i]
            try:
                result = execute_step(step, context)
            except Exception as e:
                logging.error(f"Error executing step '{step}': {e}")
                st.error("Brain down, try again shortly!")
                return None
            completed_steps.append((step, result))
            context += f"\nStep: {step}\nResult: {result}\n"

            # Update plan display to show completed steps (with checkmark)
            plan_lines = []
            for idx, s in enumerate(steps):
                if idx < len(completed_steps):
                    plan_lines.append(f"✅ **Step {idx+1}:** {s}\n")
                else:
                    plan_lines.append(f"**Step {idx+1}:** {s}\n")
            sidebar_steps.markdown(
                "\n".join(
                    [
                        (
                            f"✅ {idx+1}. {s}\n"
                            if idx < len(completed_steps)
                            else f"{idx+1}. {s}"
                        )
                        for idx, s in enumerate(steps)
                    ]
                )
            )

            # Replanning
            if not replan_limit_reached:
                try:
                    steps, replan_rounds, replan_limit_reached = replanner(
                        context, steps, replan_rounds, max_replan_rounds, replan_limit_reached
                    )
                except Exception as e:
                    logging.error(f"Error during replanning: {e}")
                    st.error("Brain down, try again shortly!")
                    return None
            i += 1  # Always increment to move to the next step

        # Final report
        try:
            report = report_writer(context)
        except Exception as e:
            logging.error(f"Error generating report: {e}")
            st.error("Brain down, try again shortly!")
            return None
        return report
    except Exception as e:
        logging.critical(f"Critical error in agentic_brain: {e}")
        st.error("Brain down, try again shortly!")
        return None

def add_hyperlink(paragraph, url, text):
    """Add a clickable hyperlink to a python-docx paragraph."""
    # This is a workaround since python-docx does not support hyperlinks natively
    # Reference: https://github.com/python-openxml/python-docx/issues/74
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Style: blue and underlined
    color = OxmlElement('w:color')
    color.set(qn('w:val'), "0000FF")
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), "single")
    rPr.append(u)

    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph

def generate_word_doc_from_markdown(markdown_text):
    try:
        # Convert markdown to HTML
        html = md.markdown(markdown_text, extensions=['tables'])
        soup = BeautifulSoup(html, "html.parser")
        doc = Document()
        doc.add_heading("DeepQuest Research Report", 0)

        for element in soup.children:
            if element.name and element.name.startswith("h") and element.name[1:].isdigit():
                level = int(element.name[1:])
                doc.add_heading(element.get_text(), level=level)
            elif element.name == "ul":
                for li in element.find_all("li", recursive=False):
                    p = doc.add_paragraph(style="List Bullet")
                    for child in li.children:
                        if getattr(child, "name", None) == "a":
                            add_hyperlink(p, child.get("href"), child.get_text())
                        elif getattr(child, "name", None) is None:
                            p.add_run(str(child))
            elif element.name == "ol":
                for li in element.find_all("li", recursive=False):
                    p = doc.add_paragraph(style="List Number")
                    for child in li.children:
                        if getattr(child, "name", None) == "a":
                            add_hyperlink(p, child.get("href"), child.get_text())
                        elif getattr(child, "name", None) is None:
                            p.add_run(str(child))
            elif element.name == "p":
                p = doc.add_paragraph()
                for child in element.children:
                    if getattr(child, "name", None) == "a":
                        add_hyperlink(p, child.get("href"), child.get_text())
                    elif getattr(child, "name", None) is None:
                        p.add_run(str(child))
            elif element.name == "table":
                rows = element.find_all("tr")
                if not rows:
                    continue
                cols = rows[0].find_all(["td", "th"])
                n_cols = len(cols)
                n_rows = len(rows)
                table = doc.add_table(rows=n_rows, cols=n_cols)
                for row_idx, row in enumerate(rows):
                    cells = row.find_all(["td", "th"])
                    for col_idx, cell in enumerate(cells):
                        cell_text = ""
                        for child in cell.children:
                            if getattr(child, "name", None) == "a":
                                cell_text += f"{child.get_text()} ({child.get('href')})"
                            elif getattr(child, "name", None) is None:
                                cell_text += str(child)
                        table.cell(row_idx, col_idx).text = cell_text
            # You can add more HTML tag handling as needed

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        logging.error(f"Error converting markdown to Word: {e}")
        return None
    
query = st.chat_input("Enter your research query:")
    
if query:
    with st.spinner("Running agentic research..."):
        try:
            report = agentic_brain(query)
            if report is None:
                st.error("Brain down, try again shortly!")
            else:
                st.subheader("Final Research Report")
                st.markdown(report)
                word_buffer = generate_word_doc_from_markdown(report)
                if word_buffer:
                    st.download_button(
                        label="Download Report as Word Document",
                        data=word_buffer,
                        file_name="deepquest_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                else:
                    st.error("Brain down, try again shortly!")
        except Exception as e:
            logging.critical(f"Critical error in main UI: {e}")
            st.error("Brain down, try again shortly!")